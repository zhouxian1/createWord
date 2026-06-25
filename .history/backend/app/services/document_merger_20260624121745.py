"""文档合稿与排版服务"""
import os
import json
import logging
from datetime import datetime
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from app import db
from app.models.models import Document as DocModel, DocumentChapter, MergeTask, Project

logger = logging.getLogger(__name__)


class DocumentMerger:
    """文档合稿服务 - 支持多文档合并与自动排版"""

    # 438C标准排版参数
    FORMAT_CONFIG = {
        'page': {
            'size': 'A4',
            'orientation': 'portrait',
            'top_margin': 2.54,   # cm
            'bottom_margin': 2.54,
            'left_margin': 3.17,
            'right_margin': 3.17
        },
        'fonts': {
            'title': {'name': '黑体', 'size': 22, 'bold': True},
            'heading1': {'name': '黑体', 'size': 16, 'bold': True},
            'heading2': {'name': '黑体', 'size': 14, 'bold': True},
            'heading3': {'name': '黑体', 'size': 12, 'bold': True},
            'heading4': {'name': '楷体', 'size': 12, 'bold': True},
            'body': {'name': '宋体', 'size': 12, 'bold': False},
            'table_header': {'name': '黑体', 'size': 10, 'bold': True},
            'table_body': {'name': '宋体', 'size': 10, 'bold': False}
        },
        'paragraph': {
            'line_spacing': 1.5,
            'space_before': 0,
            'space_after': 6,
            'first_line_indent': 2  # 首行缩进2字符
        }
    }

    def merge_documents(self, task_id: int) -> MergeTask:
        """合并多个文档"""
        task = MergeTask.query.get(task_id)
        if not task:
            raise ValueError(f"合稿任务不存在: {task_id}")

        task.status = 'running'
        db.session.commit()

        try:
            source_ids = json.loads(task.source_documents)
            documents = DocModel.query.filter(DocModel.id.in_(source_ids)).all()

            if not documents:
                raise ValueError("没有找到源文档")

            if task.merge_strategy == 'sequential':
                output_path = self._merge_sequential(documents, task)
            elif task.merge_strategy == 'by_type':
                output_path = self._merge_by_type(documents, task)
            else:
                output_path = self._merge_sequential(documents, task)

            task.status = 'completed'
            task.output_path = output_path
            task.completed_at = datetime.utcnow()
        except Exception as e:
            task.status = 'failed'
            task.error_message = str(e)
            logger.error(f"文档合稿失败: {str(e)}")
        finally:
            db.session.commit()

        return task

    def _merge_sequential(self, documents: List[DocModel], task: MergeTask) -> str:
        """顺序合并文档"""
        doc = self._create_formatted_document()

        for i, document in enumerate(documents):
            if i > 0:
                doc.add_page_break()

            # 添加文档标题
            self._add_title(doc, document.doc_name)

            # 添加各章节内容
            chapters = DocumentChapter.query.filter_by(document_id=document.id)\
                .order_by(DocumentChapter.sort_order).all()

            for chapter in chapters:
                self._add_chapter(doc, chapter)

        return self._save_document(doc, task)

    def _merge_by_type(self, documents: List[DocModel], task: MergeTask) -> str:
        """按文档类型合并"""
        doc = self._create_formatted_document()

        # 按类型分组
        type_groups = {}
        for d in documents:
            type_groups.setdefault(d.doc_type, []).append(d)

        for doc_type, docs in type_groups.items():
            for document in docs:
                doc.add_page_break()
                self._add_title(doc, document.doc_name)
                chapters = DocumentChapter.query.filter_by(document_id=document.id)\
                    .order_by(DocumentChapter.sort_order).all()
                for chapter in chapters:
                    self._add_chapter(doc, chapter)

        return self._save_document(doc, task)

    def generate_document_file(self, document_id: int) -> str:
        """将单个文档生成为DOCX文件"""
        document = DocModel.query.get(document_id)
        if not document:
            raise ValueError(f"文档不存在: {document_id}")

        doc = self._create_formatted_document()

        # 添加封面信息
        self._add_cover(doc, document)

        # 添加各章节
        chapters = DocumentChapter.query.filter_by(document_id=document_id)\
            .order_by(DocumentChapter.sort_order).all()

        for chapter in chapters:
            self._add_chapter(doc, chapter)

        # 保存文件
        output_dir = os.path.join(os.path.dirname(__file__), '..', 'static', 'uploads', 'generated')
        os.makedirs(output_dir, exist_ok=True)
        filename = f"{document.doc_code}_{document.doc_name}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)

        # 更新文档记录
        document.file_path = output_path
        db.session.commit()

        return output_path

    def _create_formatted_document(self) -> Document:
        """创建格式化的文档"""
        doc = Document()

        # 设置页面格式
        section = doc.sections[0]
        config = self.FORMAT_CONFIG['page']
        section.page_width = Cm(21.0)  # A4
        section.page_height = Cm(29.7)
        section.top_margin = Cm(config['top_margin'])
        section.bottom_margin = Cm(config['bottom_margin'])
        section.left_margin = Cm(config['left_margin'])
        section.right_margin = Cm(config['right_margin'])

        return doc

    def _add_cover(self, doc: Document, document: DocModel):
        """添加封面"""
        project = document.project

        # 添加空行
        for _ in range(6):
            doc.add_paragraph('')

        # 文档标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(document.doc_name)
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 系统信息
        for _ in range(3):
            doc.add_paragraph('')

        info_items = [
            f"系统名称：{project.system_name if project else ''}",
            f"系统版本：{project.system_version if project else ''}",
            f"文档编号：{document.doc_code or ''}",
            f"编制单位：{project.organization if project else ''}",
            f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}"
        ]

        for item in info_items:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(item)
            run.font.size = Pt(14)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        doc.add_page_break()

    def _add_title(self, doc: Document, title: str):
        """添加标题"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        config = self.FORMAT_CONFIG['fonts']['title']
        run = para.add_run(title)
        run.font.size = Pt(config['size'])
        run.font.bold = config['bold']
        run.font.name = config['name']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), config['name'])

    def _add_chapter(self, doc: Document, chapter: DocumentChapter):
        """添加章节内容"""
        # 根据层级选择标题样式
        level_key = f'heading{chapter.chapter_level}' if chapter.chapter_level <= 4 else 'heading4'
        font_config = self.FORMAT_CONFIG['fonts'].get(level_key, self.FORMAT_CONFIG['fonts']['body'])

        # 添加章节标题
        heading = doc.add_heading(level=chapter.chapter_level)
        run = heading.add_run(f"{chapter.chapter_number}  {chapter.chapter_title}")
        run.font.size = Pt(font_config['size'])
        run.font.bold = font_config['bold']
        run.font.name = font_config['name']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_config['name'])

        # 添加章节内容
        if chapter.content:
            # 处理内容中的段落
            paragraphs = chapter.content.split('\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                para = doc.add_paragraph()
                body_config = self.FORMAT_CONFIG['fonts']['body']
                run = para.add_run(para_text)
                run.font.size = Pt(body_config['size'])
                run.font.name = body_config['name']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), body_config['name'])

                # 设置段落格式
                para_format = self.FORMAT_CONFIG['paragraph']
                para.paragraph_format.line_spacing = para_format['line_spacing']
                para.paragraph_format.space_before = Pt(para_format['space_before'])
                para.paragraph_format.space_after = Pt(para_format['space_after'])
                if para_text[0] not in ('-', '*', '1', '2', '3', '•'):
                    para.paragraph_format.first_line_indent = Pt(body_config['size'] * para_format['first_line_indent'])

    def _save_document(self, doc: Document, task: MergeTask) -> str:
        """保存文档"""
        output_dir = os.path.join(os.path.dirname(__file__), '..', 'static', 'uploads', 'generated')
        os.makedirs(output_dir, exist_ok=True)
        filename = f"merged_{task.name}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)
        return output_path
