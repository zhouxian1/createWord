"""文档合稿与排版服务 - 支持多文档合并、自动排版、页眉页脚、目录生成"""
import os
import json
import logging
from datetime import datetime
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

from app import db
from app.models.models import Document as DocModel, DocumentChapter, MergeTask, Project

logger = logging.getLogger(__name__)


class DocumentMerger:
    """文档合稿服务 - 支持多文档合并与自动排版"""

    # 438C标准排版参数 (GJB排版规范)
    FORMAT_CONFIG = {
        'page': {
            'size': 'A4',
            'orientation': 'portrait',
            'top_margin': 2.54,   # cm
            'bottom_margin': 2.54,
            'left_margin': 3.17,
            'right_margin': 3.17
        },
        'fonts': {
            'title': {'name': '黑体', 'size': 22, 'bold': True},
            'heading1': {'name': '黑体', 'size': 16, 'bold': True},
            'heading2': {'name': '黑体', 'size': 14, 'bold': True},
            'heading3': {'name': '黑体', 'size': 12, 'bold': True},
            'heading4': {'name': '楷体', 'size': 12, 'bold': True},
            'body': {'name': '宋体', 'size': 12, 'bold': False},
            'table_header': {'name': '黑体', 'size': 10, 'bold': True},
            'table_body': {'name': '宋体', 'size': 10, 'bold': False},
            'header': {'name': '宋体', 'size': 9, 'bold': False},
            'footer': {'name': '宋体', 'size': 9, 'bold': False},
            'toc': {'name': '宋体', 'size': 12, 'bold': False},
        },
        'paragraph': {
            'line_spacing': 1.5,
            'space_before': 0,
            'space_after': 6,
            'first_line_indent': 2  # 首行缩进2字符
        },
        'header_footer': {
            'header_margin': 1.5,  # cm
            'footer_margin': 1.75,
            'header_distance': 1.5,
            'footer_distance': 1.75,
        }
    }

    def merge_documents(self, task_id: int) -> MergeTask:
        """合并多个文档"""
        task = MergeTask.query.get(task_id)
        if not task:
            raise ValueError(f"合稿任务不存在: {task_id}")

        task.status = 'running'
        db.session.commit()

        try:
            source_ids = json.loads(task.source_documents)
            documents = DocModel.query.filter(DocModel.id.in_(source_ids)).all()

            if not documents:
                raise ValueError("没有找到源文档")

            if task.merge_strategy == 'sequential':
                output_path = self._merge_sequential(documents, task)
            elif task.merge_strategy == 'by_type':
                output_path = self._merge_by_type(documents, task)
            else:
                output_path = self._merge_sequential(documents, task)

            task.status = 'completed'
            task.output_path = output_path
            task.completed_at = datetime.utcnow()
        except Exception as e:
            task.status = 'failed'
            task.error_message = str(e)
            logger.error(f"文档合稿失败: {str(e)}")
        finally:
            db.session.commit()

        return task

    def _merge_sequential(self, documents: List[DocModel], task: MergeTask) -> str:
        """顺序合并文档"""
        doc = self._create_formatted_document()

        # 设置页眉页脚
        self._setup_header_footer(doc, task.header_text, task.footer_text, documents[0] if documents else None)

        # 添加封面
        if task.include_cover and documents:
            self._add_cover(doc, documents[0])

        # 添加目录
        if task.include_toc:
            self._add_toc(doc)

        # 按顺序添加各文档内容
        for i, document in enumerate(documents):
            if i > 0 or task.include_cover:
                doc.add_page_break()

            # 添加文档标题
            self._add_title(doc, document.doc_name)

            # 添加各章节内容
            chapters = DocumentChapter.query.filter_by(document_id=document.id)\
                .order_by(DocumentChapter.sort_order).all()

            for chapter in chapters:
                self._add_chapter(doc, chapter)

            # 文档间页眉页脚衔接
            if i < len(documents) - 1:
                self._add_section_break(doc, documents[i + 1], task)

        return self._save_document(doc, task)

    def _merge_by_type(self, documents: List[DocModel], task: MergeTask) -> str:
        """按文档类型合并"""
        doc = self._create_formatted_document()

        # 设置页眉页脚
        self._setup_header_footer(doc, task.header_text, task.footer_text, documents[0] if documents else None)

        # 添加封面
        if task.include_cover and documents:
            self._add_cover(doc, documents[0])

        # 添加目录
        if task.include_toc:
            self._add_toc(doc)

        # 按类型分组
        type_groups = {}
        for d in documents:
            type_groups.setdefault(d.doc_type, []).append(d)

        for doc_type, docs in type_groups.items():
            for document in docs:
                doc.add_page_break()
                self._add_title(doc, document.doc_name)
                chapters = DocumentChapter.query.filter_by(document_id=document.id)\
                    .order_by(DocumentChapter.sort_order).all()
                for chapter in chapters:
                    self._add_chapter(doc, chapter)

        return self._save_document(doc, task)

    def generate_document_file(self, document_id: int) -> str:
        """将单个文档生成为DOCX文件（存储至MinIO）"""
        document = DocModel.query.get(document_id)
        if not document:
            raise ValueError(f"文档不存在: {document_id}")

        doc = self._create_formatted_document()

        # 设置页眉页脚
        self._setup_header_footer(doc, None, None, document)

        # 添加封面信息
        self._add_cover(doc, document)

        # 添加目录
        self._add_toc(doc)

        # 添加各章节
        chapters = DocumentChapter.query.filter_by(document_id=document_id)\
            .order_by(DocumentChapter.sort_order).all()

        for chapter in chapters:
            self._add_chapter(doc, chapter)

        # 保存到临时文件
        import tempfile
        import uuid
        filename = f"{document.doc_code}_{document.doc_name}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        doc.save(tmp_path)

        # 上传到MinIO
        from app.services.storage_service import get_storage
        storage = get_storage()
        project_id = document.project_id or 'default'
        object_name = f"generated/{project_id}/{uuid.uuid4().hex}.docx"
        upload_result = storage.upload_file(object_name, tmp_path)

        # 清理临时文件
        os.unlink(tmp_path)

        # 更新文档记录（存储MinIO对象名）
        document.file_path = object_name
        db.session.commit()

        return object_name

    def _create_formatted_document(self) -> Document:
        """创建格式化的文档"""
        doc = Document()

        # 设置页面格式
        section = doc.sections[0]
        config = self.FORMAT_CONFIG['page']
        section.page_width = Cm(21.0)  # A4
        section.page_height = Cm(29.7)
        section.top_margin = Cm(config['top_margin'])
        section.bottom_margin = Cm(config['bottom_margin'])
        section.left_margin = Cm(config['left_margin'])
        section.right_margin = Cm(config['right_margin'])

        return doc

    def _setup_header_footer(self, doc: Document, header_text: str = None,
                              footer_text: str = None, document: DocModel = None):
        """设置页眉页脚"""
        for section in doc.sections:
            # 页眉
            header = section.header
            header.is_linked_to_previous = False
            if header.paragraphs:
                header_para = header.paragraphs[0]
            else:
                header_para = header.add_paragraph()

            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_config = self.FORMAT_CONFIG['fonts']['header']

            # 构建页眉文本
            h_text = header_text or ''
            if not h_text and document:
                project = document.project
                h_text = f"{project.system_name if project else ''} - {document.doc_name}"

            run = header_para.add_run(h_text)
            run.font.size = Pt(header_config['size'])
            run.font.name = header_config['name']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), header_config['name'])

            # 页眉下方添加分隔线
            pPr = header_para._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

            # 页脚
            footer = section.footer
            footer.is_linked_to_previous = False
            if footer.paragraphs:
                footer_para = footer.paragraphs[0]
            else:
                footer_para = footer.add_paragraph()

            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_config = self.FORMAT_CONFIG['fonts']['footer']

            # 构建页脚文本（含页码）
            f_text = footer_text or ''
            run = footer_para.add_run(f_text + ' ' if f_text else '')
            run.font.size = Pt(footer_config['size'])
            run.font.name = footer_config['name']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), footer_config['name'])

            # 添加页码字段
            fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run1 = footer_para.add_run()
            run1._r.append(fldChar1)

            instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
            run2 = footer_para.add_run()
            run2._r.append(instrText)

            fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run3 = footer_para.add_run()
            run3._r.append(fldChar2)

    def _add_section_break(self, doc: Document, next_document: DocModel, task: MergeTask):
        """添加分节符（用于文档间页眉页脚衔接）"""
        # 添加新节
        new_section = doc.add_section()
        new_section.page_width = Cm(21.0)
        new_section.page_height = Cm(29.7)
        config = self.FORMAT_CONFIG['page']
        new_section.top_margin = Cm(config['top_margin'])
        new_section.bottom_margin = Cm(config['bottom_margin'])
        new_section.left_margin = Cm(config['left_margin'])
        new_section.right_margin = Cm(config['right_margin'])

        # 设置新节的页眉页脚
        self._setup_header_footer_for_section(new_section, task.header_text, task.footer_text, next_document)

    def _setup_header_footer_for_section(self, section, header_text: str = None,
                                           footer_text: str = None, document: DocModel = None):
        """为特定节设置页眉页脚"""
        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            header_para = header.paragraphs[0]
        else:
            header_para = header.add_paragraph()

        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_config = self.FORMAT_CONFIG['fonts']['header']

        h_text = header_text or ''
        if not h_text and document:
            project = document.project
            h_text = f"{project.system_name if project else ''} - {document.doc_name}"

        run = header_para.add_run(h_text)
        run.font.size = Pt(header_config['size'])
        run.font.name = header_config['name']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), header_config['name'])

        # 页眉分隔线
        pPr = header_para._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # 页脚
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            footer_para = footer.paragraphs[0]
        else:
            footer_para = footer.add_paragraph()

        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_config = self.FORMAT_CONFIG['fonts']['footer']

        f_text = footer_text or ''
        run = footer_para.add_run(f_text + ' ' if f_text else '')
        run.font.size = Pt(footer_config['size'])
        run.font.name = footer_config['name']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), footer_config['name'])

        # 页码
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run1 = footer_para.add_run()
        run1._r.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2 = footer_para.add_run()
        run2._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3 = footer_para.add_run()
        run3._r.append(fldChar2)

    def _add_cover(self, doc: Document, document: DocModel):
        """添加封面"""
        project = document.project

        # 添加空行
        for _ in range(6):
            doc.add_paragraph('')

        # 文档标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(document.doc_name)
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 英文标题
        if hasattr(document, 'doc_code'):
            en_para = doc.add_paragraph()
            en_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            from app.config.standard_documents import STANDARD_438C_DOCUMENTS
            doc_config = STANDARD_438C_DOCUMENTS.get(document.doc_type, {})
            en_name = doc_config.get('full_name', '')
            if en_name:
                run = en_para.add_run(en_name)
                run.font.size = Pt(16)
                run.font.name = 'Times New Roman'

        # 系统信息
        for _ in range(3):
            doc.add_paragraph('')

        info_items = [
            f"系统名称：{project.system_name if project else ''}",
            f"系统版本：{project.system_version if project else ''}",
            f"文档编号：{document.doc_code or ''}",
            f"编制单位：{project.organization if project else ''}",
            f"密　　级：{project.security_level if project else '内部'}",
            f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}"
        ]

        for item in info_items:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(item)
            run.font.size = Pt(14)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        doc.add_page_break()

    def _add_toc(self, doc: Document):
        """添加目录"""
        # 目录标题
        toc_title = doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_title.add_run('目　录')
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 添加目录域代码
        para = doc.add_paragraph()
        fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run1 = para.add_run()
        run1._r.append(fldChar_begin)

        instrText = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve">'
            f' TOC \\o "1-4" \\h \\z \\u '
            f'</w:instrText>'
        )
        run2 = para.add_run()
        run2._r.append(instrText)

        fldChar_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3 = para.add_run()
        run3._r.append(fldChar_separate)

        # 占位文本
        run4 = para.add_run('（请在Word中右键点击此处，选择"更新域"以生成目录）')
        run4.font.size = Pt(10)
        run4.font.color.rgb = RGBColor(128, 128, 128)

        fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5 = para.add_run()
        run5._r.append(fldChar_end)

        doc.add_page_break()

    def _add_title(self, doc: Document, title: str):
        """添加标题"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        config = self.FORMAT_CONFIG['fonts']['title']
        run = para.add_run(title)
        run.font.size = Pt(config['size'])
        run.font.bold = config['bold']
        run.font.name = config['name']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), config['name'])

    def _add_chapter(self, doc: Document, chapter: DocumentChapter):
        """添加章节内容"""
        # 根据层级选择标题样式
        level_key = f'heading{chapter.chapter_level}' if chapter.chapter_level <= 4 else 'heading4'
        font_config = self.FORMAT_CONFIG['fonts'].get(level_key, self.FORMAT_CONFIG['fonts']['body'])

        # 添加章节标题
        heading = doc.add_heading(level=min(chapter.chapter_level, 4))
        run = heading.add_run(f"{chapter.chapter_number}  {chapter.chapter_title}")
        run.font.size = Pt(font_config['size'])
        run.font.bold = font_config['bold']
        run.font.name = font_config['name']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_config['name'])

        # 添加章节内容
        if chapter.content:
            paragraphs = chapter.content.split('\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                # 检测是否为表格行
                if '\t' in para_text and para_text.count('\t') >= 2:
                    self._add_table_row(doc, para_text)
                    continue

                para = doc.add_paragraph()
                body_config = self.FORMAT_CONFIG['fonts']['body']
                run = para.add_run(para_text)
                run.font.size = Pt(body_config['size'])
                run.font.name = body_config['name']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), body_config['name'])

                # 设置段落格式
                para_format = self.FORMAT_CONFIG['paragraph']
                para.paragraph_format.line_spacing = para_format['line_spacing']
                para.paragraph_format.space_before = Pt(para_format['space_before'])
                para.paragraph_format.space_after = Pt(para_format['space_after'])
                # 非列表项首行缩进
                if para_text[0] not in ('-', '*', '1', '2', '3', '4', '5', '6', '7', '8', '9', '•', 'a', 'b', 'c'):
                    para.paragraph_format.first_line_indent = Pt(body_config['size'] * para_format['first_line_indent'])

    def _add_table_row(self, doc: Document, row_text: str):
        """添加表格行（简化处理）"""
        cells = row_text.split('\t')
        if not cells:
            return

        # 查找最后一个表格
        tables = doc.tables
        if tables:
            last_table = tables[-1]
            # 检查列数是否匹配
            if len(last_table.columns) == len(cells):
                row = last_table.add_row()
                for i, cell_text in enumerate(cells):
                    if i < len(row.cells):
                        row.cells[i].text = cell_text.strip()
                return

        # 创建新表格
        table = doc.add_table(rows=1, cols=len(cells))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, cell_text in enumerate(cells):
            cell = table.rows[0].cells[i]
            cell.text = cell_text.strip()
            # 表头样式
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _save_document(self, doc: Document, task: MergeTask) -> str:
        """保存文档"""
        output_dir = os.path.join(os.path.dirname(__file__), '..', 'static', 'uploads', 'generated')
        os.makedirs(output_dir, exist_ok=True)
        filename = f"merged_{task.name}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)
        return output_path
